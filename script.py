import os
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_HEADER_FOOTER

input_folder = 'files'
output_folder = 'destination'
logo_path = os.path.join(input_folder, 'logo.jpeg')

os.makedirs(output_folder, exist_ok=True)

def remove_images_from_header_footer(header_footer):
    for paragraph in header_footer.paragraphs:
        for run in paragraph.runs:
            if run.element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                run.clear()

def add_image_to_footer(footer, image_path, section_width):
    for paragraph in footer.paragraphs:
        p_element = paragraph._element
        p_element.getparent().remove(p_element)
    
    paragraph = footer.add_paragraph()
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    run = paragraph.add_run()
    
    if os.path.exists(image_path):
        largura_maxima = (section_width / 914400) * 0.9
        run.add_picture(image_path, width=Inches(largura_maxima))
    else:
        print("Imagem não encontrada:", image_path)

def process_document(doc_path, new_doc_path):
    doc = Document(doc_path)

    for section in doc.sections:
        section_width = section.page_width - section.left_margin - section.right_margin
        
        remove_images_from_header_footer(section.header)
        remove_images_from_header_footer(section.footer)
        
        add_image_to_footer(section.footer, logo_path, section_width)
    
    doc.save(new_doc_path)

for filename in os.listdir(input_folder):
    if filename.endswith('.docx'):
        input_file_path = os.path.join(input_folder, filename)
        output_file_path = os.path.join(output_folder, f"{os.path.splitext(filename)[0]}_NEW.docx")
        
        print(f"Processando: {filename}")
        process_document(input_file_path, output_file_path)

print("Processamento concluído!")